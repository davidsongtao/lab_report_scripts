"""
Description: 
    
-*- Encoding: UTF-8 -*-
@File     ：utils.py
@Author   ：King Songtao
@Time     ：2024/9/11 上午7:18
@Contact  ：king.songtao@gmail.com
"""
import shutil
from datetime import datetime
from docx.shared import Pt
from log_config import *
from docx.oxml.ns import qn
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta


def create_reports(report_ids, project_names, test_type, template):
    """
    创建检测报告
    复制模板文件，并创建新的检测报告，将所有检测报告按照report_ids中给定的报告编号依次命名
    1. report_ids -> 要创建的所有报告的报告编号列表，该数据来自于委托统计
    2. output_dir -> 要输出的新报告保存的位置。该参数可以在项目总参数配置中更改
    """
    # 检查模板文件是否存在
    if not os.path.exists(template):
        logger.info(f"模板文件未找到！请检查配置文件！")
        return
    part_2 = f"{project_names[0]}_{test_type[0]}_" + datetime.now().strftime("%Y%m%d%H%M") + "_待打印"
    output_dir = os.path.join(r"E:\检测报告生成", part_2)

    # 检查输出文件路径是否存在，不存在则创建
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
            logger.info(f"输出目录{output_dir}不存在，已创建！")
        except Exception as e:
            logger.error(f"创建输出目录时发生错误！错误信息：{e}")
            return

    for i, name in enumerate(report_ids):
        try:
            # 构造新的文件路径
            new_file_path = os.path.join(output_dir, f"{name}_{project_names[i]}.docx")
            # 复制模板文件到新的路径
            shutil.copy(template, new_file_path)
            logger.info(f"检测报告 >>>{name}_{project_names[i]}.docx<<< 已成功创建！")
        except Exception as e:
            logger.error(f"创建检测报告时发生错误！错误信息：{e}")

    return output_dir


def set_font(run, font_name, font_size):
    """设置字体和字号"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def update_content(new_text, old_text, font_size=10, center=False, reports_path=param.output_directory):
    """更新报告中的内容"""
    doc = Document(reports_path)

    def replace_in_runs(runs):
        text = "".join([run.text for run in runs])
        if old_text in text:
            new_text_parts = text.replace(old_text, new_text).split()
            for i, run in enumerate(runs):
                if i < len(new_text_parts):
                    run.text = new_text_parts[i]
                    set_font(run, "宋体", font_size)
                    if center:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置居中
                else:
                    run.text = ""

    for para in doc.paragraphs:
        replace_in_runs(para.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs)

    doc.save(reports_path)


def convert_sampling_date_to_report_date(sampling_date, days_added):
    """将委托日期转换成报告日期，增加指定的天数"""
    report_date = []
    for date_str in sampling_date:
        # 将字符串日期转换为 datetime 对象
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        # 在日期上加一天
        new_date_obj = date_obj + timedelta(days=days_added)
        # 将新的日期对象转换回字符串
        new_date_str = new_date_obj.strftime('%Y-%m-%d')
        # 将新的日期字符串添加到新列表中
        report_date.append(new_date_str)
    return report_date
